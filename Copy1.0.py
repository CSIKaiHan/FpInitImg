import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_event_name_illustrations(task_folder):
    dir_name = os.path.basename(task_folder)
    split_name = dir_name.split('_')
    assert split_name[0].isdigit(), '資料夾名稱非「案號+案由」'
    event_number = split_name[0]
    event_name = f'{split_name[1]}「{split_name[2]}」'
    illustrations = []
    fin_file_list = []
    for file_name in os.listdir(task_folder):
        if file_name.endswith(('.JPG', '.jpg', '.PNG', '.png')):
            char1 = 'oxvrp'
            for i in range(len(char1)):
                if char1[i] in os.path.splitext(file_name)[0]:
                    fin_file_list.append(file_name)
                    break

    for fin_file in fin_file_list:
        number = os.path.splitext(fin_file)[0]
        char2 = 'oxp'
        for i in range(len(char2)):
            if char2[i] in number:
                number = number.replace(char2[i], '')
        char3 = 'vr'
        for i in range(len(char3)):
            if char3[i] in number:
                number = number[:number.find(char3[i])]

        filepath = os.path.join(task_folder, fin_file)

        illustration = ''
        if 'o' in os.path.splitext(fin_file)[0]:
            illustration = f'編號{number}指紋足資比對'
        elif 'x' in os.path.splitext(fin_file)[0]:
            illustration = f'編號{number}指紋不足資比對'
        elif 'v' in os.path.splitext(fin_file)[0]:
            _ = os.path.splitext(fin_file)[0][os.path.splitext(fin_file)[0].find('v') + 1:]
            illustration = f'編號{number}為被害人{_}'
        elif 'r' in os.path.splitext(fin_file)[0]:
            _ = os.path.splitext(fin_file)[0][os.path.splitext(fin_file)[0].find('r') + 1:]
            illustration = f'編號{number}為關係人{_}'
        if 'p' in os.path.splitext(fin_file)[0]:
            illustration = illustration.replace('p', '')
            illustration = illustration.replace('指', '掌')
        illustrations.append([number, filepath, illustration])
        illustrations = sort_hierarchical_codes(illustrations)
    return event_number, event_name, illustrations


def make_docx(is_table, event_number, event_name, illustrations):
    doc1 = Document('init1.docx')
    section = doc1.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(f'{event_name}指紋初鑑照片')
    header_run.bold = True
    header_run.font.name = '標楷體'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    header_run.font.size = Pt(18)

    first_paragraph = doc1.paragraphs[0]
    new_paragraph = first_paragraph.insert_paragraph_before()

    for number, filepath, description in illustrations:
        new_paragraph.add_run().add_picture(filepath, width=Inches(2.9))
        run = new_paragraph.add_run(description)
        run.bold = True
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(14)
    if len(illustrations) % 4 == 1 or len(illustrations) % 4 == 3:
        for i in range(17):
            temp_paragraph = doc1.add_paragraph()
            doc1.paragraphs[0]._p.addnext(temp_paragraph._p)

    doc1.save(f'{os.path.join(task_folder, event_number)}指紋初鑑照片.docx')

    if is_table:
        doc2 = Document('init2.docx')
        table_number = []
        for number, filepath, description in illustrations:
            if '紋足' in description:
                table_number.append(number)

        first_row = doc2.tables[0].rows[0]
        first_cell = first_row.cells[0]
        paragraph = first_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rewrite = paragraph.text
        paragraph.clear()
        run = paragraph.add_run(rewrite.replace('113○○○○○○○', event_number))
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(14)

        for i, row_data in enumerate(table_number):
            row_index = 2 + i
            while row_index + 1 >= len(doc2.tables[0].rows):
                doc2.tables[0].add_row()

            # 處理第一欄（序號）
            cell = doc2.tables[0].cell(row_index, 0)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.clear()
            run = paragraph.add_run(str(i + 1))
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(12)

            # 處理第二欄（編號）
            cell = doc2.tables[0].cell(row_index, 1)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.clear()
            run = paragraph.add_run(str(row_data))
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(12)

        cell = doc2.tables[0].cell(len(table_number) + 2, 1)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        run = paragraph.add_run('以下空白')
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(12)

        doc2.save(f'{os.path.join(task_folder, event_number)}指紋初鑑結果彙整表.docx')


def natural_sort_key(code):
    code = code[0]
    code = os.path.splitext(code)[0]
    parts = code.split('-')
    result = []

    for part in parts:
        if part.isdigit():
            result.append((0, int(part)))
        else:
            result.append((1, part))

    MAX_DEPTH = 4
    while len(result) < MAX_DEPTH:
        result.append((0, 0))

    return tuple(result)


def sort_hierarchical_codes(codes):
    return sorted(codes, key=natural_sort_key)


if __name__ == '__main__':
    intro = '****** 歡迎使用 初鑑資料產生器1.0 ******     by TNPD phoenixkai'
    print(intro)

    print('#選擇任務資料夾')
    root = tk.Tk()
    root.withdraw()
    task_folder = filedialog.askdirectory()
    print(task_folder)

    event_number, event_name, illustrations = get_event_name_illustrations(task_folder)
    while True:
        input_value = input('#是否需要初鑑結果彙整表?<y/n>')
        if input_value == 'y' or input_value == 'Y':
            is_table = True
            break
        elif input_value == 'n' or input_value == 'N':
            is_table = False
            break

    make_docx(is_table, event_number, event_name, illustrations)
