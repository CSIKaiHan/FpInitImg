import os
import re
import tkinter as tk
import sys
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


def get_event_name_illustrations(task_folder):
    dir_name = os.path.basename(task_folder)
    split_name = dir_name.split('_')
    event_number = split_name[0]
    event_name = f'{split_name[1]}「{split_name[2]}」'
    illustrations = []
    fin_file_list = []
    for file_name in os.listdir(task_folder):
        if file_name.endswith(('.JPG', '.jpg', '.PNG', '.png')):
            char1 = 'oxvrp'
            for i in range(len(char1)):
                if char1[i] in os.path.splitext(file_name)[0]:
                    fin_file_list.append(file_name)
                    break

    for fin_file in fin_file_list:
        number = os.path.splitext(fin_file)[0]
        char2 = 'oxp'
        for i in range(len(char2)):
            if char2[i] in number:
                number = number.replace(char2[i], '')
        char3 = 'vr'
        for i in range(len(char3)):
            if char3[i] in number:
                number = number[:number.find(char3[i])]

        filepath = os.path.join(task_folder, fin_file)

        illustration = ''
        if 'o' in os.path.splitext(fin_file)[0]:
            illustration = f'編號{number}指紋足資比對'
        elif 'x' in os.path.splitext(fin_file)[0]:
            illustration = f'編號{number}指紋不足資比對'
        elif 'v' in os.path.splitext(fin_file)[0]:
            _ = os.path.splitext(fin_file)[0][os.path.splitext(fin_file)[0].find('v') + 1:]
            illustration = f'編號{number}為被害人{_}'
        elif 'r' in os.path.splitext(fin_file)[0]:
            _ = os.path.splitext(fin_file)[0][os.path.splitext(fin_file)[0].find('r') + 1:]
            illustration = f'編號{number}為關係人{_}'
        if 'p' in os.path.splitext(fin_file)[0]:
            illustration = illustration.replace('p', '')
            illustration = illustration.replace('指', '掌')
        illustrations.append([number, filepath, illustration])
        illustrations = sort_hierarchical_codes(illustrations)
    return event_number, event_name, illustrations


def make_img_docx(event_number, event_name, illustrations):
    doc1 = Document('init1.docx')
    section = doc1.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(f'{event_name}指紋初鑑照片')
    header_run.bold = True
    header_run.font.name = '標楷體'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    header_run.font.size = Pt(18)

    count = 0
    for number, filepath, description in illustrations:
        ratio = adjust_image(filepath)
        if ratio < 2/3:
            doc1.paragraphs[count].add_run().add_picture(filepath, height=Inches(4.35))
        else:
            doc1.paragraphs[count].add_run().add_picture(filepath, width=Inches(2.9))
        temp_paragraph = doc1.add_paragraph()
        doc1.paragraphs[count]._p.addnext(temp_paragraph._p)
        count += 1
        run = doc1.paragraphs[count].add_run(description)
        run.bold = True
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(14)
        if count < 2 * len(illustrations) - 1:
            temp_paragraph = doc1.add_paragraph()
            doc1.paragraphs[count]._p.addnext(temp_paragraph._p)
        count += 1
    if len(illustrations) % 4 == 1 or len(illustrations) % 4 == 3:
        for i in range(18):
            temp_paragraph = doc1.add_paragraph()
            doc1.paragraphs[count-1]._p.addnext(temp_paragraph._p)
    doc1.save(f'{os.path.join(task_folder, event_number)}指紋初鑑照片.docx')


def make_table_docx(event_number, illustrations):
    doc2 = Document('init2.docx')
    table_number = []
    for number, filepath, description in illustrations:
        if '紋足' in description:
            table_number.append(number)

    first_row = doc2.tables[0].rows[0]
    first_cell = first_row.cells[0]
    paragraph = first_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rewrite = paragraph.text
    paragraph.clear()
    run = paragraph.add_run(rewrite.replace('113○○○○○○○', event_number))
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(14)

    for i, row_data in enumerate(table_number):
        row_index = 2 + i
        while row_index + 1 >= len(doc2.tables[0].rows):
            doc2.tables[0].add_row()

        cell = doc2.tables[0].cell(row_index, 0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        run = paragraph.add_run(str(i + 1))
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(12)

        cell = doc2.tables[0].cell(row_index, 1)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        run = paragraph.add_run(str(row_data))
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(12)

    cell = doc2.tables[0].cell(len(table_number) + 2, 1)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    run = paragraph.add_run('以下空白')
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(12)
    doc2.save(f'{os.path.join(task_folder, event_number)}指紋初鑑結果彙整表.docx')

def make_table_xlsx(event_number, event_name, illustrations):
    wb = load_workbook('init3.xlsx')
    ws = wb.active

    illustration = []
    for number, filepath, _ in illustrations:
        if 'o' in os.path.splitext(filepath)[0]:
            illustration.append((number,'未比中','','否'))
        elif 'x' in os.path.splitext(filepath)[0]:
            illustration.append((number,'無法比對','','否'))
        elif 'v' in os.path.splitext(filepath)[0]:
            _ = os.path.splitext(filepath)[0][os.path.splitext(filepath)[0].find('v') + 1:]
            illustration.append((number,'比中',f'被害人{_}','否'))
        elif 'r' in os.path.splitext(filepath)[0]:
            _ = os.path.splitext(filepath)[0][os.path.splitext(filepath)[0].find('r') + 1:]
            illustration.append((number,'比中',f'關係人{_}','否'))

    start_row = 10

    if len(illustration) > 6:
        ws.delete_rows(start_row + len(illustration), 50 - len(illustration))
    else:
        ws.delete_rows(start_row + len(illustration), 44)

    for i, value in enumerate(illustration):
        ws.cell(row=start_row + i, column=1, value=i + 1)
        for j in range(4):
            ws.cell(row=start_row + i, column=j+2, value=value[j])

    cell_d1 = ws['D1']
    cell_d1.value = cell_d1.value.replace('#', str(event_number))

    cell_a3 = ws['A3']
    cell_a3.value = cell_a3.value.replace('#', event_name)

    img = XLImage('頁尾.png')
    img.width = 743
    img.height = 454
    ws.add_image(img, 'A'+str(ws.max_row))

    wb.save(f'{os.path.join(task_folder, event_number)}初鑑報告書.xlsx')


def natural_sort_key(code):
    code = code[0]
    code = os.path.splitext(code)[0]

    parts = re.findall(r'([A-Za-z]+|\d+)', code)
    result = []

    for part in parts:
        if part.isdigit():
            result.append((0, int(part)))
        else:
            result.append((1, part))

    MAX_DEPTH = 4
    while len(result) < MAX_DEPTH:
        result.append((0, 0))

    return tuple(result)


def sort_hierarchical_codes(codes):
    return sorted(codes, key=natural_sort_key)

def adjust_image(input_path, dpi=(300, 300)):
    with Image.open(input_path) as img:
        cleaned_img = Image.new(img.mode, img.size)
        cleaned_img.putdata(list(img.getdata()))
        width, height = cleaned_img.size

        if width > height:
            processed_img = cleaned_img.rotate(-90, expand=True)
        else:
            processed_img = cleaned_img
        ratio = processed_img.width / processed_img.height
        original_dpi = img.info.get('dpi', None)

        if original_dpi is None:
            processed_img.save(input_path, format=img.format, dpi=dpi, quality=100)
        else:
            processed_img.save(input_path, format=img.format, dpi=original_dpi, quality=100)

        return ratio


def main():
    intro = '****** 歡迎使用 初鑑資料產生器1.0 ******     by TNPD phoenixkai'
    print(intro)

    print('#選擇任務資料夾')
    root = tk.Tk()
    root.withdraw()
    global task_folder
    task_folder = filedialog.askdirectory()
    print(task_folder)
    try:
        event_number, event_name, illustrations = get_event_name_illustrations(task_folder)
    except Exception:
        messagebox.showerror('錯誤', '資料夾名稱非「案號_○○分局_案由」')
        sys.exit(1)

    while True:
        print('#需要什麼表? 0-不需要、1-初鑑結果彙整表、2-初鑑報告書')
        input_value = input()
        if input_value == '0' or input_value == '1' or input_value == '2':
            break
    try:
        make_img_docx(event_number, event_name, illustrations)
        if input_value == '0':
            pass
        elif input_value == '1':
            make_table_docx(event_number, illustrations)
        elif input_value == '2':
            make_table_xlsx(event_number, event_name, illustrations)
        messagebox.showinfo('成功', '執行成功！')
    except Exception as e:
        print(e)
        messagebox.showerror('錯誤', '執行失敗！')
        sys.exit(1)

task_folder = ''

if __name__ == '__main__':
    main()